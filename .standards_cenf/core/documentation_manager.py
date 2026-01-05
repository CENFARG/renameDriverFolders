# src/core/documentation_manager.py
import logging
from pathlib import Path
from typing import Any, Dict, Optional, Union

from src.core.file_manager import FileManager # Asumiendo que FileManager está en el mismo core module

logger = logging.getLogger(__name__)

class DocumentationManager:
    """
    Gestiona la generación de documentación en diferentes formatos (ej. Mintlify Markdown, DOCX).
    """

    def __init__(self, file_manager: FileManager, docs_base_path: Optional[Union[str, Path]] = None):
        """
        Inicializa el DocumentationManager.
        Args:
            file_manager: Una instancia de FileManager para manejar las operaciones de archivo.
            docs_base_path: Ruta base donde se guardará la documentación generada.
                            Si no se proporciona, se usará una ruta por defecto.
        """
        self.file_manager = file_manager
        self.docs_base_path = Path(docs_base_path) if docs_base_path else self.file_manager.get_path(relative_path="docs_output")
        self.file_manager.ensure_directory(self.docs_base_path)
        logger.info(f"DocumentationManager inicializado. Ruta base de docs: {self.docs_base_path}")

    def _generate_markdown_content(self, data: Dict[str, Any], template_name: Optional[str] = None) -> str:
        """
        Genera contenido Markdown a partir de un diccionario de datos.
        Puede usar una plantilla simple o un formato por defecto.
        """
        content = f"# {data.get("title", "Documento sin título")}\n\n"
        content += f"**Fecha:** {data.get("date", "N/A")}\n\n"
        content += f"**Descripción:** {data.get("description", "Sin descripción.")}\n\n"

        if "sections" in data and isinstance(data["sections"], list):
            for section in data["sections"]:
                content += f"## {section.get("heading", "")}\n\n"
                content += f"{section.get("body", "")}\n\n"
                if "items" in section and isinstance(section["items"], list):
                    for item in section["items"]:
                        content += f"- {item}\n"
                    content += "\n"
        
        # Ejemplo de cómo se podría integrar con Mintlify (requiere estructura específica)
        # Para Mintlify, los archivos suelen tener un frontmatter YAML
        frontmatter = f"""---
title: {data.get("title", "Documento sin título")}
description: {data.get("description", "Sin descripción.")}
---

"""
        return frontmatter + content

    def generate_mintlify_doc(self, doc_name: str, data: Dict[str, Any], subdirectory: Optional[Union[str, Path]] = None) -> Path:
        """
        Genera un archivo Markdown compatible con Mintlify.
        Args:
            doc_name: Nombre del archivo (sin extensión).
            data: Diccionario de datos para poblar el documento.
            subdirectory: Subdirectorio dentro de docs_base_path para guardar el documento.
        Returns:
            La ruta al archivo Markdown generado.
        """
        target_dir = self.docs_base_path
        if subdirectory:
            target_dir = target_dir / subdirectory
            self.file_manager.ensure_directory(target_dir)

        file_path = target_dir / f"{doc_name}.md"
        markdown_content = self._generate_markdown_content(data)
        self.file_manager.write_text_file(file_path, markdown_content)
        logger.info(f"Documento Mintlify generado: {file_path}")
        return file_path

    def generate_docx_doc(self, doc_name: str, data: Dict[str, Any], subdirectory: Optional[Union[str, Path]] = None) -> Path:
        """
        Genera un archivo DOCX a partir de un diccionario de datos.
        Requiere la librería `python-docx`.
        (Implementación placeholder - requiere lógica específica de python-docx)
        Args:
            doc_name: Nombre del archivo (sin extensión).
            data: Diccionario de datos para poblar el documento.
            subdirectory: Subdirectorio dentro de docs_base_path para guardar el documento.
        Returns:
            La ruta al archivo DOCX generado.
        """
        try:
            from docx import Document
            from docx.shared import Inches

            document = Document()
            document.add_heading(data.get("title", "Documento DOCX"), level=1)
            document.add_paragraph(f"Fecha: {data.get("date", "N/A")}")
            document.add_paragraph(data.get("description", ""))

            if "sections" in data and isinstance(data["sections"], list):
                for section in data["sections"]:
                    document.add_heading(section.get("heading", ""), level=2)
                    document.add_paragraph(section.get("body", ""))
                    if "items" in section and isinstance(section["items"], list):
                        for item in section["items"]:
                            document.add_paragraph(f"- {item}")

            target_dir = self.docs_base_path
            if subdirectory:
                target_dir = target_dir / subdirectory
                self.file_manager.ensure_directory(target_dir)

            file_path = target_dir / f"{doc_name}.docx"
            document.save(file_path)
            logger.info(f"Documento DOCX generado: {file_path}")
            return file_path
        except ImportError:
            logger.info("La librería 'python-docx' no está instalada. No se puede generar DOCX.")
            # En un escenario real, podrías querer lanzar una excepción o manejarlo de otra forma.
            # raise # Descomentar si se desea que falle explícitamente
            return None # O retornar None si la generación falla por falta de dependencia
        except Exception as e:
            logger.error(f"Error al generar documento DOCX: {e}")
            raise

# Example usage (for testing purposes, not part of the class itself)
if __name__ == "__main__":
    import os
    import tempfile
    import shutil
    import json
    from src.core.file_manager import FileManager # Asumiendo que FileManager está en el mismo core module

    # Configurar un logger básico para la prueba
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    
    # Crear un directorio temporal para las pruebas
    temp_test_dir = Path(tempfile.mkdtemp())
    print(f"Directorio temporal para pruebas: {temp_test_dir}")

    # Crear un config.template.json temporal para la prueba de ConfigManager
    # Asegurarse de que el directorio 'config' exista
    temp_config_dir = Path(__file__).parent.parent.parent / "config"
    temp_config_dir.mkdir(parents=True, exist_ok=True)
    temp_config_path = temp_config_dir / "config.template.json"
    config_data = {
        "app_info": {"name": "TestDocApp", "version": "1.0.0"},
        "logging": {"level": "DEBUG", "file": "logs/test_doc_app.log"},
        "paths": {
            "docs_output_dir": "docs_generated"
        }
    }
    with open(temp_config_path, 'w', encoding='utf-8') as f:
        json.dump(config_data, f, indent=4)

    try:
        # Inicializar ConfigManager y LoggerManager
        # Asumiendo que ConfigManager y LoggerManager están correctamente implementados y accesibles
        # Necesitamos simular su inicialización si no están disponibles en este scope
        class MockConfigManager:
            def __init__(self, config_path):
                self.config_path = config_path
                self.config = self._load_config()
            
            def _load_config(self):
                with open(self.config_path, 'r', encoding='utf-8') as f:
                    return json.load(f)

            def get_setting(self, key, default=None):
                keys = key.split('.')
                value = self.config
                try:
                    for k in keys:
                        value = value[k]
                    return value
                except (KeyError, TypeError):
                    return default

        class MockLoggerManager:
            @staticmethod
            def initialize(config_manager):
                # Simula la inicialización del logger
                log_level_str = config_manager.get_setting('logging.level', 'INFO')
                log_level = getattr(logging, log_level_str.upper(), logging.INFO)
                log_file = config_manager.get_setting('logging.file', 'app.log')
                
                # Configurar el logger básico si no está ya configurado
                if not logging.getLogger().handlers:
                    logging.basicConfig(level=log_level, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
                    # Opcionalmente, añadir un FileHandler si se desea loguear a archivo
                    # file_handler = logging.FileHandler(log_file)
                    # logging.getLogger().addHandler(file_handler)
                logger.info("MockLoggerManager initialized.")

        # Simular FileManager
        class MockFileManager:
            def __init__(self, base_path, config_manager):
                self.base_path = Path(base_path)
                self.config_manager = config_manager
                self.logger = logging.getLogger(__name__)

            def get_path(self, relative_path: Optional[str] = None, config_key: Optional[str] = None) -> Path:
                if config_key:
                    path_from_config = self.config_manager.get_setting(config_key)
                    if path_from_config:
                        return self.base_path / path_from_config
                if relative_path:
                    return self.base_path / relative_path
                return self.base_path

            def ensure_directory(self, path: Union[str, Path]):
                p = Path(path)
                p.mkdir(parents=True, exist_ok=True)
                self.logger.debug(f"Ensured directory: {p}")

            def write_text_file(self, file_path: Union[str, Path], content: str):
                p = Path(file_path)
                self.ensure_directory(p.parent)
                with open(p, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.logger.debug(f"Wrote to file: {p}")

        config_manager = MockConfigManager(config_path=temp_config_path)
        MockLoggerManager.initialize(config_manager)

        file_manager = MockFileManager(base_path=temp_test_dir, config_manager=config_manager)

        # Inicializar DocumentationManager
        doc_manager = DocumentationManager(file_manager=file_manager)

        # Datos de ejemplo para el documento
        doc_data = {
            "title": "Reporte de Prueba",
            "date": "2025-08-26",
            "description": "Este es un reporte generado automáticamente para fines de prueba.",
            "sections": [
                {
                    "heading": "Sección 1: Introducción",
                    "body": "Aquí va el contenido de la primera sección. Puede ser texto largo."
                },
                {
                    "heading": "Sección 2: Detalles",
                    "body": "Más detalles sobre el tema.",
                    "items": [
                        "Item A: Descripción del item A",
                        "Item B: Descripción del item B",
                        "Item C: Descripción del item C"
                    ]
                }
            ]
        }

        # Generar documento Mintlify (Markdown)
        mintlify_doc_path = doc_manager.generate_mintlify_doc("mi_reporte_mintlify", doc_data, subdirectory="reports")
        print(f"Documento Mintlify generado en: {mintlify_doc_path}")

        # Generar documento DOCX (requiere python-docx instalado)
        try:
            docx_doc_path = doc_manager.generate_docx_doc("mi_reporte_docx", doc_data, subdirectory="reports")
            if docx_doc_path:
                print(f"Documento DOCX generado en: {docx_doc_path}")
        except ImportError:
            print("Skipping DOCX generation: python-docx not installed.")
        except Exception as e:
            print(f"An error occurred during DOCX generation: {e}")

    finally:
        # Limpiar el directorio temporal
        if temp_test_dir.exists():
            shutil.rmtree(temp_test_dir)
            print(f"Directorio temporal {temp_test_dir} eliminado.")
        # Limpiar el config.template.json temporal
        if temp_config_path.exists():
            temp_config_path.unlink()
        # Limpiar el directorio de config si está vacío
        try:
            temp_config_dir.rmdir()
        except OSError:
            # El directorio puede no estar vacío si hay otros archivos de config
            pass
